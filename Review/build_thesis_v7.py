from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parent
REFERENCE = ROOT / "KLTN_InterV_LeMinhDuy_v6.docx"
OUTPUT = ROOT / "KLTN_InterV_LeMinhDuy_v7.docx"


def set_paragraph_text(paragraph, text: str) -> None:
    """Replace visible text while retaining paragraph and first-run formatting."""
    paragraph_element = paragraph._p
    first_run_properties = None
    for run in paragraph.runs:
        if run._r.rPr is not None:
            first_run_properties = deepcopy(run._r.rPr)
            break

    for child in list(paragraph_element):
        if child.tag != qn("w:pPr"):
            paragraph_element.remove(child)

    run_element = OxmlElement("w:r")
    if first_run_properties is not None:
        run_element.append(first_run_properties)
    text_element = OxmlElement("w:t")
    if text[:1].isspace() or text[-1:].isspace():
        text_element.set(qn("xml:space"), "preserve")
    text_element.text = text
    run_element.append(text_element)
    paragraph_element.append(run_element)


def set_cell_text(cell, text: str) -> None:
    if not cell.paragraphs:
        cell.add_paragraph()
    set_paragraph_text(cell.paragraphs[0], text)
    for extra in cell.paragraphs[1:]:
        set_paragraph_text(extra, "")


def all_body_paragraphs(document):
    paragraphs = list(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    return paragraphs


def replace_everywhere(document, replacements: dict[str, str]) -> None:
    for paragraph in all_body_paragraphs(document):
        original = paragraph.text
        updated = original
        for old, new in replacements.items():
            updated = updated.replace(old, new)
        if updated != original:
            set_paragraph_text(paragraph, updated)


def clone_row_after(table, row_index: int, values: list[str]) -> None:
    template = deepcopy(table.rows[row_index]._tr)
    table._tbl.insert(row_index + 1, template)
    new_row = table.rows[row_index + 1]
    for index, value in enumerate(values):
        set_cell_text(new_row.cells[index], value)


def append_appendix(document) -> None:
    document.add_page_break()
    document.add_paragraph(
        "PHỤ LỤC E. ĐỐI CHIẾU SNAPSHOT SOURCE V7",
        style="Heading 1",
    )
    document.add_paragraph(
        "Phụ lục này ghi lại các điểm đã đối chiếu khi cập nhật khóa luận từ bản v6 sang v7. Snapshot dùng để viết bản v7 là commit e7fdfb9 ngày 19 tháng 08 năm 2026. Các con số được lấy từ kiểm kê source và kết quả chạy kiểm thử, không phải số liệu ước lượng.",
        style="Normal",
    )
    document.add_paragraph(
        "E.1. Khác biệt về quy mô và hợp đồng",
        style="Heading 2",
    )
    document.add_paragraph(
        "So với snapshot trước, mã nguồn hiện có 499 tệp văn bản và 72.615 dòng; 61 tệp route Next.js cung cấp 73 phương thức HTTP; 39 trang ứng dụng, 19 mô hình Mongoose, 16 RPC gRPC và 86 tệp Markdown trong corpus rule. Hai RPC bổ sung là GenerateOpeningTurn và ExtractCandidateProfile. Đây là cơ sở để cập nhật lại các bảng kiểm kê, sơ đồ contract và phần phụ lục RPC.",
        style="Normal",
    )
    document.add_paragraph(
        "E.2. Luồng hội thoại sau khi bổ sung phần giới thiệu",
        style="Heading 2",
    )
    document.add_paragraph(
        "Phiên practice hiện không chuyển thẳng từ TTS sang câu hỏi chuyên môn. Sau khi ứng viên trả lời lời giới thiệu, route opening lưu transcript và metadata, gRPC GenerateOpeningTurn gọi DeepSeek với JD, rule và evidence, sau đó trả về acknowledgement, transition, câu hỏi đầu tiên, spoken_text và transition_type. Backend kiểm tra schema, grounding ID và độ dài trước khi lưu câu hỏi q_1. Các lượt sau dùng ba kiểu chuyển tiếp continue_competency, probe_gap và bridge_to_next_competency; opening_to_first chỉ được phép ở lượt mở đầu.",
        style="Normal",
    )
    document.add_paragraph(
        "E.3. Speech, realtime và khả năng phục hồi",
        style="Heading 2",
    )
    document.add_paragraph(
        "Recorder realtime có hai chế độ: phiên AssemblyAI mới cho từng lượt hoặc persistent session cho auto-turn-taking. AudioWorklet gửi PCM vào WebSocket, MediaRecorder lưu bản ghi, còn RealtimeTranscript lọc final turn trùng và late message theo turn order hoặc segment generation. Khi transcript realtime rỗng, hệ thống dùng Faster-Whisper; khi microphone hoặc streaming không khả dụng, người dùng có thể chuyển sang trả lời bằng text nếu phiên cho phép. Các trạng thái UI openingSpeaking, openingRecording, reviewing, submitting và finishing được tách riêng để tránh gửi lặp hoặc dùng callback của lượt cũ.",
        style="Normal",
    )
    document.add_paragraph(
        "E.4. Dữ liệu kết quả và chi phí vận hành",
        style="Heading 2",
    )
    document.add_paragraph(
        "PracticeRun lưu candidateIntro, transition metadata, audio reference, provider transcript và usage snapshot. Luồng finish bắt buộc phải có audio để phân tích SenseVoice; với interview tuyển dụng, CandidateProfile được trích xuất song song với evaluation và chỉ giữ các mục có evidence từ transcript, đồng thời loại bỏ nhóm trường nhạy cảm. Voice preview được hash theo language, voice và sample text, kiểm tra MongoDB trước khi gọi TTS, lưu audioBase64 dưới dạng bản ghi dùng chung và có khóa chống request trùng. PaymentSettlement bổ sung pending recovery, check-status, cancel, reconcile PayOS và settlement idempotent trong MongoDB transaction.",
        style="Normal",
    )
    document.add_paragraph(
        "E.5. Bằng chứng kiểm chứng và giới hạn diễn giải",
        style="Heading 2",
    )
    document.add_paragraph(
        "Snapshot hiện tại đạt 51/51 kiểm thử backend, 6/6 kiểm thử RealtimeTranscript, lint frontend và production build Next.js. Kiểm tra TypeScript độc lập bằng tsc --noEmit còn báo TS5097 tại import có phần mở rộng .ts trong frontend/tests/RealtimeTranscript.test.ts; lỗi này không xuất hiện trong bước kiểm tra TypeScript của next build. Vì vậy bản v7 ghi nhận build production đạt nhưng không tuyên bố tsc độc lập hoàn toàn sạch. Chưa có browser E2E với microphone, WebSocket và provider thật trong lần rà soát này; các kết luận về trải nghiệm realtime vẫn phải được xem là kết quả từ code, unit/contract test và protocol thiết kế.",
        style="Normal",
    )


def main() -> None:
    document = Document(REFERENCE)

    replace_everywhere(
        document,
        {
            "468 tệp văn bản": "499 tệp văn bản",
            "72.445 dòng": "72.615 dòng",
            "54 tệp route Next.js": "61 tệp route Next.js",
            "67 phương thức HTTP": "73 phương thức HTTP",
            "18 mô hình Mongoose": "19 mô hình Mongoose",
            "14 RPC gRPC": "16 RPC gRPC",
            "14 RPC hiện tại": "16 RPC hiện tại",
            "14 RPC": "16 RPC",
            "38 kiểm thử backend": "51 kiểm thử backend",
            "Danh sách 14 RPC": "Danh sách 16 RPC",
        },
    )

    paragraphs = list(document.paragraphs)
    replacements_by_index = {
        35: "Khóa luận trình bày việc xây dựng InterV, một hệ thống phỏng vấn và luyện tập phỏng vấn tích hợp trí tuệ nhân tạo cho ứng viên, nhà tuyển dụng và quản trị viên. Nền tảng hỗ trợ hai luồng practice và recruitment, trong đó phiên practice có phần giới thiệu ứng viên, câu hỏi có grounding, follow-up thích ứng, realtime transcript và kết quả luyện tập; phiên recruitment bổ sung JD, campaign, invitation, hồ sơ ứng viên và bước xem xét của con người.",
        36: "Đóng góp kỹ thuật trọng tâm là pipeline hội thoại có cấu trúc. DeepSeek đảm nhiệm trích xuất JD, sinh câu hỏi, tạo opening turn từ phần giới thiệu, tạo follow-up và tổng hợp evaluation theo JSON schema. Qdrant cung cấp evidence với grounding ID được allow-list; SenseVoice cung cấp tín hiệu âm thanh quan sát được; AssemblyAI và Faster-Whisper tạo lớp transcript dự phòng. Hệ thống cũng bổ sung auto-turn-taking, text-only fallback, lưu preview TTS dùng chung, lịch sử practice phân trang và đối soát PayOS theo trạng thái.",
        37: "Snapshot source được kiểm kê gồm 499 tệp văn bản với 72.615 dòng; 61 tệp route Next.js cung cấp 73 phương thức HTTP, 39 trang ứng dụng, 19 mô hình Mongoose, 16 RPC gRPC và 86 tệp Markdown trong corpus rule. Kết quả xác minh hiện có 51 kiểm thử backend và 6 kiểm thử RealtimeTranscript đạt; lint và production build frontend đạt. Kiểm tra tsc độc lập còn một lỗi cấu hình import trong file test, còn browser E2E với microphone và WebSocket chưa được thực hiện trong lần rà soát này.",
        41: "This thesis presents InterV, an AI-assisted interview and interview-practice system for candidates, recruiters, and administrators. The platform supports practice and recruitment workflows. The practice flow now includes a candidate introduction, a grounded opening turn, adaptive follow-up questions, realtime transcription, optional automatic turn-taking, text fallback, and a persisted analysis result. The recruitment flow adds job descriptions, campaigns, invitations, candidate profile evidence, and human review.",
        42: "The main technical contribution is a structured conversation pipeline. DeepSeek performs job-description extraction, question generation, opening-turn generation, follow-up generation, and schema-constrained evaluation. Qdrant supplies allow-listed evidence identifiers; AssemblyAI and Faster-Whisper provide transcript paths; SenseVoice produces observable delivery signals without treating emotion tags as psychological or hiring decisions. The current source snapshot contains 499 text files, 61 Next.js route files, 73 HTTP methods, 19 Mongoose models, 16 gRPC RPCs, and 86 rule Markdown files. Backend and realtime tests, lint, and production build evidence are reported with their exact verification limits.",
        62: "Mục tiêu tổng quát là xây dựng và đánh giá về mặt kỹ thuật một hệ thống phỏng vấn và luyện tập phỏng vấn tích hợp AI, trong đó kết quả sinh và đánh giá được ràng buộc bởi JD, rule, bằng chứng truy hồi và trạng thái phiên. Mục tiêu cụ thể gồm: hỗ trợ practice và recruitment trong cùng kiến trúc nhưng tách quyền và dữ liệu; nối phần giới thiệu ứng viên với competency đầu tiên bằng opening turn có kiểm chứng; duy trì hội thoại realtime có fallback; lưu kết quả, usage và audit có thể truy nguyên; và xác định rõ giới hạn trước khi dùng AI trong quyết định tuyển dụng.",
        70: "Đối tượng và phạm vi: Đối tượng nghiên cứu gồm quy trình phỏng vấn có cấu trúc, cơ chế tạo opening turn và follow-up, đánh giá có grounding, xử lý tiếng nói đa nhà cung cấp, phân quyền theo vai trò, dữ liệu tuyển dụng, thanh toán và quan sát hệ thống. Phạm vi mã nguồn được kiểm kê gồm 499 tệp văn bản với 72.615 dòng; 61 tệp route Next.js cung cấp 73 phương thức HTTP, 39 trang ứng dụng, 19 mô hình Mongoose, 16 RPC gRPC và 86 tệp Markdown trong corpus rule.",
        80: "Một kiến trúc end-to-end cho hai chế độ practice và recruitment, có opening turn sau phần giới thiệu, follow-up theo evidence gap, frontend/BFF, gRPC AI, MongoDB giao dịch, event store và Qdrant.",
        81: "Một pipeline DeepSeek có schema JSON, grounding allow-list, repair có giới hạn, transition type cho hội thoại và telemetry usage/cost theo từng operation.",
        82: "Một pipeline tiếng nói phân vai giữa AssemblyAI streaming, Faster-Whisper fallback, Edge TTS và SenseVoice; hỗ trợ persistent session, auto-turn-taking, text-only fallback và bất biến không suy luận tâm lý từ emotion tag.",
        83: "Một corpus 86 rule được gắn provenance, cùng cơ chế lưu candidate profile chỉ từ transcript evidence và loại bỏ các nhóm trường nhạy cảm.",
        84: "Một lớp vận hành bổ sung gồm voice preview cache trong MongoDB, lịch sử practice phân trang, đối soát PayOS nhiều trạng thái và settlement credit idempotent.",
        99: "Trong từng lượt, câu hỏi được đọc bằng TTS và hiển thị bằng văn bản. Ở phần mở đầu, người dùng nghe opening prompt, ghi âm hoặc nhập text, sau đó transcript được lưu qua route opening. Backend gọi GenerateOpeningTurn để nối phần giới thiệu với competency đầu tiên bằng acknowledgement, transition và câu hỏi q_1. Ở các lượt tiếp theo, recorder gửi audio và transcript realtime; SubmitAnswer kiểm tra ownership, trạng thái, thứ tự câu và idempotency trước khi lưu QA, sau đó nhận follow-up có transition hoặc lấy câu hỏi đã được lookahead chuẩn bị. Auto-turn-taking có thể tự chốt final turn, còn text-only mode cho phép tiếp tục khi microphone hoặc streaming không dùng được.",
        100: "Kết thúc phiên không đồng nghĩa kết quả luôn tồn tại. Hệ thống chỉ gọi EvaluateInterview khi QA history hợp lệ và SenseVoice đã phân tích ít nhất một audio sample. Với recruitment interview, finish route có thể chạy CandidateProfile extraction song song với evaluation; profile chỉ giữ thông tin có evidence từ transcript. Nếu audio không có, transcript rỗng hoặc JSON evaluation không vượt qua schema/grounding gate, run phải lưu trạng thái lỗi có thể giải thích thay vì nhận điểm mặc định giả.",
        101: "Kịch bản tuyển dụng: recruiter xây dựng campaign và mô tả công việc trước khi thêm ứng viên. JD có thể được nhập trực tiếp hoặc trích xuất từ tài liệu qua ExtractJd. Sau khi kiểm tra title, response, lịch và quyền sở hữu, hệ thống tạo invitation và practice session riêng cho ứng viên. Ứng viên nhận lời mời, xác thực danh tính, trả lời opening prompt và tham gia chuỗi câu hỏi có grounding trong context của campaign.",
        102: "Khi kết thúc, backend phân tích audio delivery và đánh giá QA history. CandidateProfile extraction chạy khi có candidate introduction, trả về các mục experience, skills, education, achievements, motivation, availability, language hoặc other nếu có evidence; nhóm thông tin nhạy cảm bị loại bỏ. Recruiter xem kết quả, profile evidence và lịch sử, nhưng hệ thống không tự tuyển hoặc loại ứng viên.",
        114: "Nhóm thứ tư liên quan đến AI. Một câu hỏi phải có text, competency, difficulty, expected signals và grounding ID hợp lệ. Opening turn bắt buộc dùng transition_type=opening_to_first; follow-up chỉ dùng continue_competency, probe_gap hoặc bridge_to_next_competency. Evaluation phải dựa trên QA history của đúng run, không lấy dữ liệu phiên khác. SenseVoice analysis phải có audio sample và provider đúng hợp đồng; emotion tag không được thay đổi confidence/composure hay quyết định tuyển dụng. Các quy tắc được kiểm tra ở schema, service validation, route guard, contract test và grounding test.",
        118: "Khả năng truy nguyên của InterV được hiểu là có thể đi từ kết quả hiển thị về dữ liệu, rule, model và thao tác đã tạo ra nó. Một evaluation nên gắn run ID, phiên bản JD, profile, difficulty, QA history, candidate introduction, transition type, audio analysis provider, grounding IDs, model configuration, usage snapshot và timestamp. Candidate profile chỉ hợp lệ khi mỗi mục có evidence từ transcript; voice preview có cache key và sample hash; payment có orderCode, paymentLinkId, providerStatus và lastReconciledAt.",
        130: "Tính khả thi kỹ thuật được hỗ trợ bởi việc tách frontend/BFF và AI backend. Next.js xử lý giao diện, session, route theo vai trò, payment state và realtime recorder; Python xử lý DeepSeek, Qdrant và speech. gRPC cung cấp contract typed giữa hai runtime với 16 RPC. Các provider ngoài có API chính thức, còn Qdrant và MongoDB có thể chạy cục bộ cho phát triển. Tuy nhiên, SenseVoice yêu cầu checkpoint và tài nguyên tính toán; độ trễ realtime và tính ổn định microphone vẫn cần browser E2E trên môi trường mục tiêu.",
        141: "Mỗi câu hỏi nghiên cứu được gắn với hiện vật có thể kiểm tra để tránh kết luận dựa trên cảm nhận. RQ1 được chứng minh bằng route, model, sequence và test quyền. RQ2 dùng code RAG/DeepSeek, opening transition, follow-up grounding invariant và dữ liệu output. RQ3 dùng audio pipeline, persistent turn boundary, parser tag và invariant test. RQ4 dùng generator, manifest, candidate profile provenance và voice cache. RQ5 dùng kết quả kỹ thuật hiện có cùng danh sách bằng chứng còn thiếu, trong đó browser E2E và nghiên cứu validity/fairness vẫn là khoảng trống.",
        180: "DeepSeek trong InterV đảm nhiệm năm nhóm tác vụ: trích xuất JD, tạo bộ câu hỏi khởi đầu, tạo opening turn sau phần giới thiệu, tạo follow-up sau mỗi câu trả lời và tổng hợp đánh giá cuối phiên. Hệ thống còn dùng cùng hạ tầng để trích xuất candidate profile có lọc trường nhạy cảm. Hai cấu hình model được tách thành fast model và evaluation model; mọi output quan trọng đều đi qua schema validation, grounding allow-list và repair tối đa một lần.",
        196: "Tiền xử lý và tổng hợp nhiều đoạn: Âm thanh được nhận theo chunk cùng run ID, question ID, transcript, content type, duration và cờ final. Persistent session giữ microphone, AudioWorklet và AssemblyAI WebSocket giữa các lượt auto-turn-taking; segment ID và turn order ngăn final message cũ đi vào lượt mới. Dịch vụ chuẩn hóa mẫu thành AudioSample rồi phân tích ngoài event loop bằng asyncio.to_thread. Nếu không có đoạn âm thanh đã ghi, đánh giá SenseVoice bị từ chối thay vì tạo dữ liệu mặc định giả.",
        253: "Đối với generate questions, prompt phải yêu cầu số lượng, phân bố competency, difficulty và expected signals. Đối với opening turn, prompt phải nối introduction với competency đầu tiên bằng transition trung tính, không suy diễn hoặc lặp opening prompt. Đối với follow-up, prompt chỉ được dùng ba transition type hợp lệ, không lặp history và phải có evidence gap. Đối với evaluation và candidate profile, output phải có evidence đúng nguồn, không chứa trường nhạy cảm hoặc suy luận tâm lý. Mỗi operation dùng schema riêng để giảm ambiguity.",
        376: "Biên gRPC làm giảm ghép nối giữa TypeScript và Python. Hợp đồng proto định nghĩa 16 RPC: Health, ExtractJd, ListVoices, SynthesizeTts, StartInterview, GenerateOpeningTurn, TranscribeAudio, SubmitAnswer, EvaluateInterview, ExtractCandidateProfile, CreateStreamingToken, AnalyzeInterview, GetRagStatus, SearchKnowledge, DeleteKnowledge và GetDeepSeekBalance. InterviewTurnResponse biểu diễn câu hỏi kế tiếp cùng acknowledgement, transition, spoken_text, transition_type và usage. Mỗi response có success và payload có kiểu; thông tin bí mật không thuộc schema.",
        415: "Lớp HTTP phục vụ nhu cầu theo trang và vai trò, còn gRPC biểu diễn năng lực AI độc lập. Cách phân chia này cho phép frontend thay đổi mà không buộc dịch vụ AI biết cookie/session web; BFF chịu trách nhiệm chuyển identity đã xác thực thành request tối thiểu. 16 RPC hiện tại bao phủ health, JD, TTS, opening turn, interview, candidate profile, audio, RAG và DeepSeek balance.",
        457: "Luồng practice có các nhóm trạng thái UI: chuẩn bị, opening speaking/recording/reviewing/submitting, câu hỏi speaking/recording/reviewing/submitting, closing và finishing. Auto-turn-taking tự chốt final turn sau end_of_turn; textAnswerEnabled cho phép chuyển sang text-only. Mỗi trạng thái đồng bộ với server và có generation guard để callback của opening hoặc question cũ không ghi đè lượt mới. Nút gửi bị khóa khi answer đang commit; retry giữ khóa logic; recorder và transcript streaming được cleanup khi đổi trang hoặc kết thúc.",
        464: "Thiết kế BFF và 73 phương thức HTTP: 61 tệp route cung cấp 73 phương thức HTTP và đóng vai trò Backend for Frontend. Route auth xử lý login, logout, refresh, sessions và identity hiện tại. Route user xử lý hồ sơ, username, email, password, credit history. Route practice/recruiter điều phối nghiệp vụ và lịch sử. Route AI chuyển request sang gRPC, gồm opening, answer, finish, result, voice preview và streaming token. Route payment xử lý create, pending, check-status, verify, webhook và cancel. Route admin cung cấp quản trị và quan sát.",
        483: "Hệ thống hiện có 19 mô hình Mongoose. Ngoài user, session, run, audio, campaign, invitation và evaluation, các mô hình mới hoặc được mở rộng gồm VoicePreviewAudio cho audioBase64 dùng chung, AiUsageEvent cho từng logical operation, Transaction với providerStatus và reconciliation fields, cùng candidateIntro.items trong PracticeRun/PracticeSession. Quan hệ dùng tham chiếu cho thực thể có vòng đời độc lập và snapshot khi cần giữ ngữ cảnh tại thời điểm phỏng vấn.",
        486: "Hợp đồng gRPC ở mức thông điệp: Proto định nghĩa contract typed giữa BFF TypeScript và Python. InterviewContext chứa session_id, title, industry, job_description, topic, difficulty, question_count, language và voice_id. InterviewQuestion chứa id, text, competency, difficulty, expected_signals, tts_text và grounding_ids. InterviewTurnResponse thêm acknowledgement_text, transition_text, spoken_text và transition_type. QaPair liên kết question_id, question, answer và grounding_ids; CandidateProfileItem liên kết category, label, value và evidence.",
        492: "Module deepseek.py chia trách nhiệm thành cấu hình, usage context, HTTP client/completion, grounding preparation, generation, opening turn, follow-up, evaluation, candidate profile và balance. validate_deepseek_configuration chạy khi khởi động để không chờ đến phiên đầu mới phát hiện thiếu key/model. Mỗi completion chọn timeout fast hoặc eval. Output opening/follow-up/evaluation/profile đều được parse và kiểm tra; repair chỉ chạy một lần rồi fail closed nếu vẫn sai.",
        528: "Payment settlement tách status nội bộ khỏi providerStatus của PayOS. create route không tạo giao dịch mới khi còn pending; pending và check-status có thể khôi phục link cũ hoặc reconcile provider; cancel chỉ thành công khi PayOS xác nhận terminal state. Webhook xác minh chữ ký, orderCode, amount và paymentLinkId. Khi PAID, settlePaidTransaction cập nhật User, CreditLog và Transaction trong một MongoDB transaction, sau đó phát sự kiện credit updated; retry webhook cho kết quả already-settled và không cộng credit lần hai.",
        558: "Một deployment tối thiểu gồm Next.js process, Python gRPC service, MongoDB chính, event MongoDB, Qdrant và kết nối provider. Reverse proxy/TLS đứng trước web; gRPC có thể chỉ mở trong mạng nội bộ. Secret được inject qua environment/secret manager. VoicePreviewAudio và PracticeAudio cần giới hạn kích thước, index và retention. Health check phân biệt liveness và readiness để không nhận phiên khi model/RAG chưa sẵn sàng; browser E2E cần xác nhận microphone, AudioWorklet, WebSocket và TTS trên môi trường triển khai thật.",
        569: "Practice interview ưu tiên nhiệm vụ chính: đọc/nghe câu hỏi, chuẩn bị, ghi âm, xem transcript, chỉnh sửa hoặc gửi. Trạng thái microphone, connection, timer, realtime transcript và upload cần rõ nhưng không gây áp lực thị giác. Auto-turn-taking giảm thao tác nhưng phải cho phép dừng, xem lại và retry; text-only fallback đảm bảo phiên không phụ thuộc tuyệt đối vào microphone. Nếu TTS không phát, câu hỏi văn bản vẫn đầy đủ; nếu transcript thiếu, route transcribe dùng Faster-Whisper và hiển thị nguồn transcript.",
        613: "Mỗi bước có điểm quan sát vận hành riêng: thời lượng audio, latency STT, fallback count, segment/turn boundary, SenseVoice provider/model, warning, DeepSeek retry, schema repair, profile extraction, voice cache hit/miss, payment reconciliation và usage. Correlation ID nối request từ Next.js qua gRPC đến provider. Log không chứa audio hoặc transcript đầy đủ theo mặc định; dữ liệu nhạy được lưu trong kho nghiệp vụ có quyền và retention rõ.",
        627: "Frontend và BFF được xây dựng trên Next.js/TypeScript. Dữ liệu nghiệp vụ dùng MongoDB qua Mongoose; sự kiện có thể tách sang MongoDB riêng. Dịch vụ AI dùng Python, gRPC, DeepSeek API, Qdrant, FunASR/SenseVoice, Faster-Whisper, AssemblyAI và Edge TTS. PayOS phục vụ thanh toán, còn persistent voice preview dùng MongoDB làm cache lâu dài. Sự kết hợp này phù hợp với web giàu tương tác và workload AI có dependency Python.",
        634: "Thanh toán. Công nghệ: PayOS và PaymentSettlement; Vai trò: Tạo order, lưu checkout/QR, webhook, polling, pending recovery, cancel, reconcile và settlement credit idempotent.",
        669: "Nguyên tắc báo cáo là chỉ ghi số đo đã thực hiện. Số lượng 499 tệp, 72.615 dòng, 61 route file, 73 HTTP methods, 39 trang, 19 model, 16 RPC, 10 test files và 86 rule files đến từ script kiểm kê tại snapshot. Kết quả 51 test backend, 6 test realtime, lint và production build đến từ lệnh chạy cụ thể. Kiểm tra tsc độc lập có lỗi TS5097 tại file test; browser E2E microphone/WebSocket/provider thật chưa có trong lần chạy này.",
        680: "Auth/user tests bao gồm register PIN hết hạn, login sai nhiều lần, refresh/revoke, đổi email/password và session list. Payment tests gồm create, verify, webhook signature, webhook replay, reconcile, pending recovery, cancel theo state và settlement idempotent. AI tests gồm opening turn, follow-up transition, candidate profile filtering, audio invariant, usage aggregation, RAG provenance và gRPC contract. RealtimeTranscript tests kiểm tra partial/final turn, duplicate, segment cũ và repeated answer; lint và production build xác nhận frontend compile được.",
        793: "Đề tài đã xây dựng và rà soát InterV theo hai trục: hệ thống phần mềm và chất lượng đánh giá. Về kỹ thuật, nền tảng kết nối giao diện web, BFF, dữ liệu MongoDB, dịch vụ AI gRPC, DeepSeek, Qdrant và chuỗi xử lý tiếng nói. Hai luồng practice và recruitment được mô hình hóa bằng trạng thái, hợp đồng và sequence riêng; practice có thêm opening turn, adaptive follow-up, auto-turn-taking và text fallback. Pipeline LLM có JSON validation, grounding allow-list, profile evidence filter và telemetry usage. Payment và voice preview cũng có persistence, reconciliation và idempotency riêng.",
        800: "SenseVoice và ASR có thể sai lệch theo ngôn ngữ, accent, thiết bị và nhiễu; chưa có browser E2E đại diện cho microphone, WebSocket và auto-turn-taking trong snapshot này. Candidate profile extraction chỉ là cấu trúc hóa transcript, không phải xác minh sự thật của CV. Voice preview cache chứng minh cơ chế lưu và tái sử dụng audio nhưng chưa phải benchmark chi phí hoặc tải production. Payment reconciliation có contract/source test nhưng cần kiểm tra provider thật trong môi trường an toàn.",
        809: "Với SenseVoice, nghiên cứu nên đánh giá riêng từng nhiệm vụ LID, AED, transcript và pacing; tuyệt đối không suy từ độ chính xác SER sang khả năng đo tâm lý. Với DeepSeek, evaluation harness cần tập câu hỏi chuẩn, opening transition, prompt injection set, grounding precision, schema failure rate, candidate profile precision và kiểm tra chuyên gia. Với realtime, cần đo auto-turn-taking latency, false turn boundary và tỷ lệ fallback. Với payment/voice cache, cần kiểm tra retry, provider outage, duplicate request và retention trước production.",
        811: "Mục tiêu kiến trúc hai luồng đã được đáp ứng ở mức thiết kế và mã nguồn: practice có session/run/audio, candidate introduction, opening turn và follow-up; recruitment có campaign/invitation, candidate profile evidence và decision tách biệt. Mục tiêu grounding được đáp ứng bằng rule catalog, Qdrant hybrid retrieval, grounding ID và hậu kiểm. Mục tiêu speech được đáp ứng bằng AssemblyAI streaming, Faster-Whisper fallback, Edge TTS và SenseVoice rich tags. Mục tiêu quản trị được hỗ trợ bởi usage event, API log, audit, payment reconciliation, voice cache và provenance.",
        822: "Với SenseVoice, kiến nghị gần nhất là tiếp tục dùng namespace AudioObservations cho languages, events, pacing, transcript quality, provider, model revision và warnings; confidence/composure không nên được diễn giải như trạng thái tâm lý. Với realtime, cần đưa segment generation và turn order vào telemetry. Với DeepSeek, opening turn, follow-up và candidate profile phải tiếp tục được đánh giá bằng evidence precision, không chỉ bằng độ trôi chảy. SER nếu giữ cho nghiên cứu phải opt-in, không hiển thị recruiter ranking và không lưu quá mục đích.",
        832: "InterV đã đạt một nền tảng kỹ thuật tương đối đầy đủ và một khung nghiên cứu thận trọng cho phỏng vấn AI. Snapshot v7 bổ sung opening turn có grounding, transition hội thoại, persistent realtime recorder, text fallback, candidate profile có evidence, voice preview MongoDB cache và PayOS reconciliation. Hệ thống có thể hỗ trợ practice và recruitment workflow, nhưng các tuyên bố về hiệu lực dự báo, công bằng và mức độ sẵn sàng cho quyết định tuyển dụng vẫn cần nghiên cứu thực nghiệm độc lập.",
        868: "A.1. Danh sách 16 RPC gRPC",
    }
    for index, text in replacements_by_index.items():
        set_paragraph_text(paragraphs[index], text)

    # Keep the reference table's existing visual treatment while adding the two
    # RPCs introduced by the current proto contract.
    rpc_table = document.tables[52]
    clone_row_after(
        rpc_table,
        5,
        [
            "6",
            "GenerateOpeningTurn",
            "Unary",
            "Nối phần giới thiệu với câu hỏi chuyên môn đầu tiên.",
        ],
    )
    clone_row_after(
        rpc_table,
        9,
        [
            "10",
            "ExtractCandidateProfile",
            "Unary",
            "Trích xuất profile có evidence từ transcript, lọc trường nhạy cảm.",
        ],
    )
    # Renumber the rows after the inserted entries.
    for row_number, row in enumerate(rpc_table.rows[1:], start=1):
        set_cell_text(row.cells[0], str(row_number))

    message_table = document.tables[23]
    clone_row_after(
        message_table,
        2,
        [
            "InterviewOpeningRequest / InterviewTurnResponse",
            "run/context/opening prompt/transcript/transition/grounding",
            "opening_to_first; prompt và transcript là candidate data không tin cậy",
            "Lặp opening, suy diễn hoặc dùng ID ngoài allow-list.",
        ],
    )
    clone_row_after(
        message_table,
        len(message_table.rows) - 1,
        [
            "CandidateProfileItem",
            "category/label/value/evidence",
            "Evidence phải lấy từ transcript; nhóm nhạy cảm bị loại",
            "Biến transcript thành profile có vẻ chắc chắn hoặc lộ PII.",
        ],
    )

    append_appendix(document)
    document.core_properties.title = (
        "Xây dựng hệ thống phỏng vấn và luyện tập phỏng vấn tích hợp trí tuệ nhân tạo - v7"
    )
    document.core_properties.subject = "Khóa luận tốt nghiệp InterV - source snapshot v7"
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
