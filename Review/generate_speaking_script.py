# -*- coding: utf-8 -*-
"""
Script to generate the comprehensive 15-Minute Thesis Defense Speaking Script and Outline
for InterV Graduation Thesis (Khoá luận tốt nghiệp - Lê Minh Duy).
Outputs a professionally styled Word document (.docx).
"""

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def set_cell_background(cell, hex_color):
    """Set the background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set cell padding (in dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin_name, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        m = OxmlElement(f'w:{margin_name}')
        m.set(qn('w:w'), str(val))
        m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    tcPr.append(tcMar)

def create_document():
    doc = docx.Document()
    
    # Page setup - A4, 2cm margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        
    # Styles
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(31, 41, 55) # Gray-800
    normal_style.paragraph_format.line_spacing = 1.2
    normal_style.paragraph_format.space_after = Pt(4)
    
    # --- COVER / HEADER ---
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(2)
    run_uni = title_p.add_run("TRƯỜNG ĐẠI HỌC VĂN HIẾN\nKHOA CÔNG NGHỆ THÔNG TIN\n")
    run_uni.font.size = Pt(12)
    run_uni.font.bold = True
    run_uni.font.color.rgb = RGBColor(30, 58, 138)
    
    title_main = doc.add_paragraph()
    title_main.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_main.paragraph_format.space_after = Pt(6)
    run_main = title_main.add_run("KỊCH BẢN THUYẾT TRÌNH BẢO VỆ KHÓA LUẬN TỐT NGHIỆP (15 PHÚT)\n")
    run_main.font.size = Pt(18)
    run_main.font.bold = True
    run_main.font.color.rgb = RGBColor(15, 23, 42)
    
    run_sub = title_main.add_run("ĐỀ TÀI: XÂY DỰNG HỆ THỐNG PHỎNG VẤN VÀ LUYỆN TẬP PHỎNG VẤN TÍCH HỢP TRÍ TUỆ NHÂN TẠO (INTERV)")
    run_sub.font.size = Pt(13)
    run_sub.font.bold = True
    run_sub.font.color.rgb = RGBColor(37, 99, 235)
    
    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_p.paragraph_format.space_after = Pt(16)
    run_info = info_p.add_run("Sinh viên thực hiện: Lê Minh Duy — MSSV: 221A011220\nGVHD: ThS. Đặng Văn Lực\nNăm học: 2026")
    run_info.font.size = Pt(11)
    run_info.font.italic = True
    run_info.font.color.rgb = RGBColor(75, 85, 99)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # Helper function for Section Heading
    def add_section_heading(text, level=1):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
        run = h.add_run(text)
        run.font.bold = True
        if level == 1:
            run.font.size = Pt(15)
            run.font.color.rgb = RGBColor(30, 58, 138)
        elif level == 2:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(37, 99, 235)
        elif level == 3:
            run.font.size = Pt(11.5)
            run.font.color.rgb = RGBColor(15, 23, 42)
        return h

    # Helper function for Callout Box
    def add_callout(text, label="LƯU Ý THUYẾT TRÌNH", bg_color="F0FDF4", border_color="16A34A"):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.columns[0].width = Inches(6.6)
        cell = table.cell(0, 0)
        set_cell_background(cell, bg_color)
        set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r_lbl = p.add_run(f"📌 {label}: ")
        r_lbl.font.bold = True
        r_lbl.font.size = Pt(10.5)
        r_lbl.font.color.rgb = RGBColor(22, 101, 52)
        
        r_txt = p.add_run(text)
        r_txt.font.size = Pt(10)
        r_txt.font.italic = True
        r_txt.font.color.rgb = RGBColor(21, 128, 61)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # --- PHẦN 1: DÀN Ý & CHIẾN LƯỢC PHÂN BỔ THỜI GIAN ---
    add_section_heading("PHẦN I. DÀN Ý & CHIẾN LƯỢC PHÂN BỔ THỜI GIAN 15 PHÚT", level=1)
    
    p_strat = doc.add_paragraph()
    p_strat.add_run("Chiến lược cốt lõi: ").font.bold = True
    p_strat.add_run(
        "Hội đồng đánh giá cao nhất khả năng làm chủ kiến trúc kỹ thuật và giải quyết các bài toán phức tạp của hệ thống. "
        "Do đó, thời gian 15 phút (900 giây) được phân bổ nghiêm ngặt theo tỉ lệ vàng: "
        "5 phút cho Tổng quan, Bối cảnh, Ranh giới, FE/BE và Kết luận; "
        "10 phút trọn vẹn dành cho Pipeline Cốt Lõi (Lookahead Engine, Multimodal Speech, Hybrid RAG Grounding, SenseVoice Analysis, Grounded Evaluation)."
    )
    
    # Table of Time Allocation
    table_time = doc.add_table(rows=1, cols=4)
    table_time.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_time.autofit = False
    
    headers = ["Phần", "Nội dung & Slide trọng tâm", "Thời gian", "Mục tiêu & Tác động"]
    widths = [Inches(1.2), Inches(2.4), Inches(1.1), Inches(1.9)]
    
    hdr_cells = table_time.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].width = widths[i]
        set_cell_background(hdr_cells[i], "1E3A8A")
        set_cell_margins(hdr_cells[i], top=100, bottom=100, left=120, right=120)
        p = hdr_cells[i].paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(255, 255, 255)
        
    time_data = [
        ("Phần 1: Mở đầu & Tổng quan\n(Slide 1 - 11)", 
         "• Giới thiệu đề tài\n• Bài toán: Ảo giác LLM & suy diễn giọng nói\n• 3 Actor & Ranh giới trách nhiệm\n• Kiến trúc tổng thể 2 tầng (Next.js BFF ↔ gRPC ↔ Python)", 
         "~ 2.5 phút\n[00:00 - 02:30]", 
         "Gây ấn tượng ban đầu về tính bài bản, ranh giới rõ ràng, không tuyên bố quá mức (Responsible AI)."),
        
        ("Phần 2: PIPELINE CỐT LÕI\n(Slide 12 - 17, 26, 27)\n★ TRỌNG TÂM ★", 
         "1. Luồng tổng thể End-to-End\n2. Lookahead Adaptive Question Engine (0ms latency, next.js after())\n3. Pipeline Giọng nói Đa tầng (AssemblyAI, Faster-Whisper, Vbee TTS)\n4. RAG Hybrid & Grounding 3 Cổng (Allow-list, 1-shot repair)\n5. SenseVoice Behavioral Analysis (LID/SER/AED, Observation only)\n6. Grounded Evaluation (Exact Answer Excerpt Check >40đ)", 
         "~ 10 phút\n[02:30 - 12:30]", 
         "CHINH PHỤC HỘI ĐỒNG:\nChứng minh chiều sâu kỹ thuật, cơ chế xử lý độ trễ, khả năng chịu lỗi và tính xác thực tuyệt đối."),
        
        ("Phần 3: Tính năng phụ & Hiện vật\n(Slide 18 - 25)", 
         "• Recruiter & Admin Workspace\n• PayOS Credit Ledger & AI Usage Tracking\n• Hiện vật: 61 routes, 19 models, 16 RPCs, 51/51 tests", 
         "~ 1.5 phút\n[12:30 - 14:00]", 
         "Minh chứng quy mô hệ sinh thái hoàn chỉnh, production-ready, zero-mock runtime."),
        
        ("Phần 4: Kết luận & Q&A\n(Slide 28 - 32)", 
         "• 4 Đóng góp chính\n• Giới hạn trung thực & 4 Rollout Gates\n• Kết luận & Tiếp nhận phản biện", 
         "~ 1.0 phút\n[14:00 - 15:00]", 
         "Khép lại đĩnh đạc, khiêm tốn học thuật, tự tin bước vào phần hỏi đáp.")
    ]
    
    for row_idx, data in enumerate(time_data):
        row = table_time.add_row()
        bg_col = "F8FAFC" if row_idx % 2 == 0 else "FFFFFF"
        if "TRỌNG TÂM" in data[0]:
            bg_col = "EFF6FF"
        for col_idx, text in enumerate(data):
            cell = row.cells[col_idx]
            cell.width = widths[col_idx]
            set_cell_background(cell, bg_col)
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(text)
            r.font.size = Pt(9.5)
            if col_idx == 0 and "TRỌNG TÂM" in text:
                r.font.bold = True
                r.font.color.rgb = RGBColor(30, 58, 138)
            elif col_idx == 2:
                r.font.bold = True
                r.font.color.rgb = RGBColor(185, 28, 28)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # --- PHẦN 2: KỊCH BẢN NÓI CHI TIẾT TỪNG SLIDE ---
    add_section_heading("PHẦN II. KỊCH BẢN NÓI CHI TIẾT TỪNG SLIDE (CHÍNH XÁC 15 PHÚT)", level=1)
    
    doc.add_paragraph(
        "Kịch bản dưới đây được thiết kế thành từng chặng thời gian cụ thể. "
        "Sinh viên đọc đúng khẩu khí, tự tin, phát âm rõ ràng, không đọc lại nguyên văn slide mà diễn giải cơ chế và quyết định kỹ thuật đằng sau."
    )

    slides_script = [
        {
            "part": "CHẶNG 1: MỞ ĐẦU, BỐI CẢNH & KIẾN TRÚC TỔNG QUAN (2.5 PHÚT)",
            "slides": [
                {
                    "slide_num": "Slide 1",
                    "title": "Trang bìa - Giới thiệu đề tài",
                    "time": "00:00 - 00:30 (30 giây)",
                    "script": (
                        "Kính thưa quý Thầy Cô trong Hội đồng chấm khóa luận tốt nghiệp. "
                        "Em tên là Lê Minh Duy, sinh viên lớp Công nghệ Thông tin. "
                        "Hôm nay, em xin phép được đại diện nhóm nghiên cứu trình bày đề tài khóa luận tốt nghiệp: "
                        "'Xây dựng hệ thống phỏng vấn và luyện tập phỏng vấn tích hợp Trí tuệ Nhân tạo - InterV', "
                        "dưới sự hướng dẫn của Thầy ThS. Đặng Văn Lực.\n\n"
                        "InterV là một nền tảng phỏng vấn giọng nói hai chế độ: Luyện tập cá nhân và Phỏng vấn tuyển dụng thực tế, "
                        "kết hợp mô hình ngôn ngữ lớn DeepSeek, hệ thống truy hồi tri thức chuẩn hóa Qdrant RAG, "
                        "cùng mô hình nhận diện giọng nói và phân tích hành vi đa phương thức SenseVoice."
                    ),
                    "note": "Chào Hội đồng to, rõ ràng, tư thế đứng thẳng, phong thái tự tin, mắt bao quát toàn bộ Hội đồng."
                },
                {
                    "slide_num": "Slide 2 - 3",
                    "title": "Lộ trình & Bài toán thực tế",
                    "time": "00:30 - 01:15 (45 giây)",
                    "script": (
                        "Kính thưa Hội đồng, khi ứng dụng AI vào quy trình phỏng vấn, các giải pháp hiện nay thường mắc phải ba cạm bẫy lớn:\n"
                        "1. LLM thiếu căn cứ (Hallucination): Câu hỏi và lời nhận xét có vẻ trôi chảy nhưng bịa đặt hoặc lệch khỏi Mô tả công việc (JD).\n"
                        "2. Thiếu tính nhất quán: Phỏng vấn tự do khiến tiêu chí đánh giá giữa các ứng viên bị xê dịch.\n"
                        "3. Suy diễn quá mức từ giọng nói: Biến các tín hiệu âm thanh đơn thuần thành kết luận tâm lý thiếu cơ sở.\n\n"
                        "Vì vậy, mục tiêu tối thượng của InterV là xây dựng một hệ thống phỏng vấn 'Grounded' - mọi câu hỏi, "
                        "mọi tiêu chuẩn đánh giá đều phải truy nguyên được nguồn gốc từ JD và bộ quy tắc chuẩn hóa, "
                        "đồng thời giữ vững nguyên tắc: AI là công cụ hỗ trợ bằng chứng, con người là người ra quyết định cuối cùng."
                    ),
                    "note": "Nhấn mạnh 3 từ khóa: Căn cứ (Grounding) - Nhất quán (Consistency) - Con người quyết định (Human-in-the-loop)."
                },
                {
                    "slide_num": "Slide 4 - 6",
                    "title": "Ranh giới trách nhiệm & Phạm vi nghiên cứu",
                    "time": "01:15 - 01:45 (30 giây)",
                    "script": (
                        "Hệ thống phân tách rạch ròi ba nhóm người dùng với ranh giới quyền hạn tuyệt đối:\n"
                        "• Ứng viên sở hữu dữ liệu cá nhân, luyện tập hoặc tham gia theo thư mời.\n"
                        "• Recruiter sở hữu quyết định tuyển dụng: thiết lập JD, xem bằng chứng và đưa ra quyết định cuối.\n"
                        "• Admin chỉ vận hành hạ tầng kỹ thuật và theo dõi tài chính, không can thiệp vào kết quả tuyển dụng.\n\n"
                        "Đặc biệt, hệ thống tuyên bố rõ: InterV KHÔNG tự động tuyển/loại (No auto-hire/reject) và "
                        "KHÔNG chẩn đoán tâm lý, đảm bảo tuyệt đối tính minh bạch và an toàn AI."
                    ),
                    "note": "Lướt nhanh nhưng chắc chắn, thể hiện đạo đức nghiên cứu và tính an toàn của hệ thống."
                },
                {
                    "slide_num": "Slide 7 - 11",
                    "title": "Kiến trúc tổng thể 2 tầng (BFF & gRPC AI Boundary)",
                    "time": "01:45 - 02:30 (45 giây)",
                    "script": (
                        "Về mặt kiến trúc, InterV được chia thành 2 tầng tách biệt với hợp đồng giao tiếp rõ ràng:\n"
                        "• Tầng Web/BFF (Next.js App Router): Quản lý phiên, xác thực RBAC, WebSocket audio streaming, "
                        "và giao dịch cơ sở dữ liệu MongoDB ReplicaSet.\n"
                        "• Tầng AI Backend (Python gRPC Service): Đóng gói toàn bộ logic tính toán AI, bao gồm DeepSeek LLM, "
                        "Qdrant Vector Database, FastEmbed Dense & BM25 Sparse, SenseVoice và Edge/Vbee TTS.\n\n"
                        "Giao tiếp giữa Next.js và Python được ràng buộc nghiêm ngặt bằng giao thức gRPC với 16 RPCs typed có xác thực API key nội bộ. "
                        "Không có bất kỳ logic AI nào bị lộ trực tiếp ra Client, và không có dữ liệu giả lập (mock) trong quá trình vận hành."
                    ),
                    "note": "Chuyển ý mượt mà: 'Sau đây, em xin phép dành trọn vẹn 10 phút tiếp theo để đi sâu vào Pipeline Cốt lõi của hệ thống'."
                }
            ]
        },
        {
            "part": "CHẶNG 2: PIPELINE CỐT LÕI CỦA HỆ THỐNG (~10 PHÚT - TRỌNG TÂM THUYẾT TRÌNH)",
            "slides": [
                {
                    "slide_num": "Slide 12",
                    "title": "Cơ chế Lookahead Adaptive Question Engine (Trái tim của độ mượt phỏng vấn)",
                    "time": "02:30 - 04:30 (2 phút - 120 giây)",
                    "script": (
                        "Kính thưa Hội đồng, thách thức lớn nhất của một cuộc phỏng vấn AI đàm thoại là ĐỘ TRỄ. "
                        "Nếu sau mỗi câu trả lời của ứng viên, hệ thống mới bắt đầu gọi DeepSeek để sinh câu hỏi tiếp theo, "
                        "ứng viên sẽ phải chờ từ 3 đến 5 giây trong im lặng - điều này phá vỡ hoàn toàn trải nghiệm tương tác tự nhiên.\n\n"
                        "InterV giải quyết triệt để bài toán này bằng Kiến trúc Lookahead Adaptive Question Engine gồm 3 bước độc đáo:\n\n"
                        "1. Bước Khởi tạo (Preparation Phase): Khi người dùng bấm bắt đầu, Backend sinh ngay một bộ câu hỏi baseline chuẩn mực "
                        "dựa trên JD và nạp sẵn âm thanh (warm TTS). Câu hỏi số 1 được trả về tức thì.\n\n"
                        "2. Bước Phản hồi thời gian thực (0ms Latency): Khi ứng viên trả lời xong câu Q_i và nhấn nộp, "
                        "Next.js BFF ngay lập tức trả về câu hỏi Q_(i+1) đã có sẵn trong bộ nhớ đệm. Ứng viên không phải chờ một mili-giây nào!\n\n"
                        "3. Bước Thích ứng nền (Background Lookahead): Cùng lúc đó, Next.js kích hoạt hàm after() chạy bất đồng bộ trong nền, "
                        "gửi toàn bộ lịch sử QA qua gRPC SubmitAnswer tới Python backend. "
                        "DeepSeek kết hợp với RAG sẽ phân tích câu trả lời vừa rồi để 'bắt bài' (probe gap/khoảng trống năng lực), "
                        "từ đó sinh ra một câu hỏi thích ứng chuyên sâu Q_(i+2) kèm lời chuyển ý tự nhiên (transition), "
                        "đồng thời kích hoạt warm TTS ngay trong nền.\n\n"
                        "Khi câu Q_(i+2) sinh xong, nó được ghi đè vào vị trí câu hỏi kế tiếp trong MongoDB. "
                        "Nếu vì sự cố mạng mà Lookahead chưa kịp xong, hệ thống sẽ tự động dùng câu hỏi baseline đã warm sẵn làm fallback an toàn. "
                        "Nhờ cơ chế này, cuộc phỏng vấn luôn thích ứng linh hoạt theo câu trả lời của ứng viên mà độ trễ tương tác bằng 0."
                    ),
                    "note": "Chỉ tay vào sơ đồ Lookahead (Slide 12): Luồng người dùng bên trên (0ms) và Luồng nền Lookahead bên dưới. Đây là điểm sáng kiến trúc lớn nhất!"
                },
                {
                    "slide_num": "Slide 15",
                    "title": "Pipeline Xử lý Giọng nói Đa tầng (Multimodal Speech Pipeline)",
                    "time": "04:30 - 06:30 (2 phút - 120 giây)",
                    "script": (
                        "Tiếp theo, em xin trình bày về Pipeline Giọng nói Đa tầng của InterV. "
                        "Hệ thống xây dựng một chu trình khép kín từ Micro ứng viên đến Loa của AI:\n\n"
                        "1. Thu âm & Realtime Streaming: Trình duyệt mở AudioWorklet, thu âm PCM 16kHz và mở kết nối WebSocket trực tiếp tới AssemblyAI. "
                        "Để bảo mật tuyệt đối, Next.js cấp một Temporary Streaming Token có thời hạn ngắn (ngăn lộ API key gốc). "
                        "Chữ chạy hiển thị realtime trên màn hình giúp ứng viên theo dõi câu trả lời của mình.\n\n"
                        "2. STT Fallback Cục bộ (Local Resilient STT): Nếu mạng chập chờn, kết nối WebSocket bị đứt hoặc transcript rỗng, "
                        "file audio gốc (được nén và gửi về server) sẽ được chuyển giao cho model Faster-Whisper chạy trực tiếp trên Python Backend. "
                        "Điều này đảm bảo phiên phỏng vấn không bao giờ bị gián đoạn vì lỗi STT từ bên thứ ba.\n\n"
                        "3. TTS Phát âm Chuẩn Kỹ thuật (Edge/Vbee TTS): Một bài toán rất thực tế trong phỏng vấn IT là việc phát âm các thuật ngữ tiếng Anh "
                        "như 'React', 'Figma', 'Kubernetes', 'CI/CD', 'Microservices' khi AI nói tiếng Việt. "
                        "InterV đã tích hợp bộ thư viện phiên âm ngữ âm (Phonetic Pronunciation Library), tự động chuẩn hóa các thuật ngữ chuyên ngành "
                        "trước khi đưa vào bộ tổng hợp giọng nói, giúp giọng đọc của AI tự nhiên, chuyên nghiệp và chuẩn xác như một chuyên gia công nghệ."
                    ),
                    "note": "Nhấn mạnh cơ chế Fallback Faster-Whisper và Thư viện phiên âm thuật ngữ IT trong TTS."
                },
                {
                    "slide_num": "Slide 13 - 14, 26",
                    "title": "Cơ chế RAG Hybrid & Grounding Allow-List (Chống ảo giác LLM)",
                    "time": "06:30 - 08:30 (2 phút - 120 giây)",
                    "script": (
                        "Trọng tâm tiếp theo là Cơ chế Kiểm soát Tri thức RAG Hybrid và Grounding Allow-list nhằm loại bỏ hoàn toàn ảo giác của DeepSeek.\n\n"
                        "1. Kho tri thức chuẩn mực (Corpus Provenance): Hệ thống nạp 86 tệp quy tắc chuẩn hóa Markdown, bao gồm 5 core framework, "
                        "15 ngành nghề, 4 cấp bậc (Junior/Middle/Senior/Lead) và 60 hồ sơ năng lực chuyên sâu (Profile specs) dựa trên tiêu chuẩn STAR và thang đo BARS.\n\n"
                        "2. Truy hồi đa phương thức (Hybrid Retrieval): Sử dụng cơ sở dữ liệu vector Qdrant kết hợp song song:\n"
                        "   • Dense Vector: FastEmbed đa ngôn ngữ (paraphrase-multilingual-mpnet) nắm bắt ngữ nghĩa sâu.\n"
                        "   • Sparse Lexical: Thuật toán BM25 nắm bắt chính xác các từ khóa chuyên ngành.\n"
                        "   • Hai luồng kết quả được hợp nhất bằng thuật toán Reciprocal Rank Fusion (RRF).\n\n"
                        "3. Cơ chế Grounding 3 Cổng (Tri-Gate Control):\n"
                        "   • Cổng trước khi sinh (Pre-generation): Chuẩn hóa context và cấp phát danh sách Evidence IDs hợp lệ.\n"
                        "   • Cổng trong khi sinh (In-generation): Ép buộc DeepSeek trả về JSON có cấu trúc nghiêm ngặt và phải trích dẫn Evidence ID làm căn cứ.\n"
                        "   • Cổng sau khi sinh (Post-generation Validation Gate): Backend hậu kiểm toàn bộ citation. "
                        "Bất biến an toàn: Nếu DeepSeek trích dẫn bất kỳ ID nào không nằm trong Evidence Allow-list, câu trả lời sẽ bị từ chối ngay lập tức. "
                        "Backend tự động kích hoạt duy nhất 1 lượt Repair Request kèm allow-list chuẩn để DeepSeek tự sửa sai. "
                        "Nhờ vậy, mọi câu hỏi và nhận xét xuất xưởng đều có căn cứ pháp lý và chuyên môn xác thực 100%."
                    ),
                    "note": "Chỉ vào sơ đồ Slide 13 & 26: Chuỗi kiểm soát Trước - Trong - Sau khi sinh. Nhắc tới thuật toán RRF và 1-shot repair."
                },
                {
                    "slide_num": "Slide 15, 27",
                    "title": "Phân tích Hành vi Giọng nói Đa phương thức (SenseVoice Speech Analysis)",
                    "time": "08:30 - 10:00 (1.5 phút - 90 giây)",
                    "script": (
                        "Sau khi ứng viên hoàn thành các câu hỏi, toàn bộ audio gốc được stream qua gRPC tới module SenseVoice để phân tích tín hiệu phi ngôn ngữ:\n\n"
                        "1. Bóc tách tín hiệu chuyên sâu:\n"
                        "   • SER (Speech Emotion Recognition): Cảm xúc giọng nói (Tự tin, Trung tính, Căng thẳng, v.v.).\n"
                        "   • LID (Language Identification): Nhận diện ngôn ngữ phát âm.\n"
                        "   • AED (Audio Event Detection): Bắt các sự kiện âm thanh như tiếng cười, tiếng ho, thở dài.\n"
                        "   • Signal Metrics: Đo lường khách quan tốc độ nói (Words Per Minute), tỷ lệ khoảng lặng (Pause ratio), "
                        "độ ổn định âm lượng (Volume stability qua độ lệch chuẩn dB) và đếm số lượng từ đệm (Filler words bilingual: 'à, ừm, like, you know').\n\n"
                        "2. Bất biến An toàn Tuyệt đối (Safety Invariants):\n"
                        "   • Observation Only: Tín hiệu âm thanh chỉ dùng làm bằng chứng mô tả hỗ trợ luyện tập, TUYỆT ĐỐI không chẩn đoán tâm lý hay tính cách.\n"
                        "   • Neutral Default: Nếu tín hiệu âm thanh không rõ ràng hoặc lỗi, điểm Composure/Confidence mặc định trả về 50 (trung tính), "
                        "không bao giờ tự ý trừ điểm vô căn cứ của ứng viên."
                    ),
                    "note": "Khẳng định: InterV không suy diễn tâm lý trái phép, tuân thủ đúng nguyên tắc Responsible AI."
                },
                {
                    "slide_num": "Slide 13, 20, 24",
                    "title": "Cơ chế Chấm điểm & Đánh giá Căn cứ (Grounded Evaluation Engine)",
                    "time": "10:00 - 11:30 (1.5 phút - 90 giây)",
                    "script": (
                        "Khâu cuối cùng của pipeline là Chấm điểm và Tổng hợp báo cáo phỏng vấn:\n\n"
                        "1. Đánh giá đa chiều 7 tiêu chí: DeepSeek tổng hợp transcript và kết quả SenseVoice để chấm 7 thang điểm độc lập: "
                        "Communication, Knowledge, Problem Solving, Confidence, JD Fit, Composure, Vocal Delivery.\n\n"
                        "2. Ràng buộc bằng chứng thực tế (Exact Answer Excerpt Rule): "
                        "Để ngăn chặn tình trạng AI nhận xét chung chung hoặc tự bịa ra lỗi, Backend áp dụng quy tắc bất biến: "
                        "Mọi câu hỏi được chấm trên 40 điểm BẮT BUỘC phải trích xuất chính xác một đoạn văn bản (exact substring) từ câu trả lời thật của ứng viên làm minh chứng.\n\n"
                        "3. Khóa phiên nguyên tử & Chống Race Condition: Sử dụng cơ chế Evaluation Lease (thời hạn 370s) và MongoDB Transaction, "
                        "đảm bảo mỗi lượt phỏng vấn chỉ được chấm đúng một lần duy nhất, ngăn ngừa hoàn toàn việc trừ credit kép hoặc ghi đè kết quả."
                    ),
                    "note": "Nhấn mạnh 'Exact Answer Excerpt Rule' - đây là bằng chứng kỹ thuật rõ nét về việc chống ảo giác."
                },
                {
                    "slide_num": "Slide 16 - 17",
                    "title": "Hợp đồng Typed gRPC & Khả năng Phục hồi Hệ thống",
                    "time": "11:30 - 12:30 (1 phút - 60 giây)",
                    "script": (
                        "Để kết thúc phần Pipeline cốt lõi, em xin nhấn mạnh tính toàn vẹn của Hợp đồng Typed:\n"
                        "Toàn bộ dữ liệu trao đổi giữa TypeScript (Next.js) và Python đều thông qua Protobuf v3. "
                        "Mọi payload trước khi lưu vào MongoDB hay render lên UI đều phải vượt qua Schema Validation của Zod và Pydantic.\n\n"
                        "Hệ thống được thiết kế với cơ chế Graceful Fallback ở mọi tầng: "
                        "RAG lỗi có Baseline rules; Streaming STT lỗi có Faster-Whisper; Lookahead chậm có Prepared Baseline Questions; "
                        "Mất mạng có Idempotent Replay. Nhờ đó, InterV đạt độ ổn định và khả năng phục hồi cực cao."
                    ),
                    "note": "Tổng kết mốc 12:30. Chuyển sang phần tổng kết tính năng phụ và hiện vật."
                }
            ]
        },
        {
            "part": "CHẶNG 3: TÍNH NĂNG MỞ RỘNG, HIỆN VẬT & KIỂM THỬ (1.5 PHÚT)",
            "slides": [
                {
                    "slide_num": "Slide 18 - 23",
                    "title": "Hệ sinh thái mở rộng: Recruiter Workspace & Admin Dashboard",
                    "time": "12:30 - 13:15 (45 giây)",
                    "script": (
                        "Bên cạnh pipeline luyện tập cá nhân, InterV cung cấp hai không gian làm việc chuyên nghiệp:\n"
                        "• Recruiter Workspace: Cho phép nhà tuyển dụng tạo chiến dịch, nhập JD chuẩn hóa, gửi thư mời phỏng vấn tự động qua email queue, "
                        "khóa phiên phỏng vấn, xem báo cáo tổng hợp kèm trích dẫn và tự tay đưa ra quyết định tuyển dụng.\n"
                        "• Admin Dashboard: Giám sát toàn diện hệ thống với tính năng theo dõi số dư DeepSeek Balance trực tiếp qua gRPC, "
                        "đối soát chi phí token theo từng request, quản trị dòng tiền với cổng thanh toán PayOS qua Credit Ledger nguyên tử, "
                        "và hệ thống ApiRequestLog theo dõi latency p95 với cơ chế tự động xóa sau 7 ngày (TTL index)."
                    ),
                    "note": "Nói nhanh, súc tích, làm nổi bật tính ứng dụng thực tế của đồ án."
                },
                {
                    "slide_num": "Slide 24 - 25",
                    "title": "Hiện vật triển khai & Kết quả Xác minh Kỹ thuật",
                    "time": "13:15 - 14:00 (45 giây)",
                    "script": (
                        "Kính thưa Hội đồng, toàn bộ hệ thống đã được triển khai hoàn chỉnh với quy mô hiện vật ấn tượng:\n"
                        "• Hơn 60.000 dòng mã nguồn logic (TypeScript, TSX, Python, Protobuf).\n"
                        "• 61 Route Files với 74 HTTP Method Handlers trên Next.js BFF.\n"
                        "• 19 Mongoose Data Models quản lý toàn diện vòng đời dữ liệu.\n"
                        "• 16 gRPC RPCs đóng gói trọn vẹn các dịch vụ AI.\n"
                        "• 86 Rule Markdown Files bao phủ 15 ngành nghề và 60 hồ sơ vị trí.\n\n"
                        "Về mặt chất lượng kiểm thử: Toàn bộ 51/51 Backend Unit & Contract Tests đạt 100% trong 3.28 giây, "
                        "vượt qua Type Check TypeScript nghiêm ngặt, cùng các bài kiểm thử E2E tích hợp từ đăng ký, thanh toán đến hoàn thành buổi phỏng vấn."
                    ),
                    "note": "Nhấn mạnh các con số chuẩn xác: 51/51 tests pass, 16 gRPC RPCs, 19 models, 61 routes."
                }
            ]
        },
        {
            "part": "CHẶNG 4: ĐÓNG GÓP, GIỚI HẠN & KẾT LUẬN (1.0 PHÚT)",
            "slides": [
                {
                    "slide_num": "Slide 28 - 30",
                    "title": "Đóng góp & Giới hạn nghiên cứu",
                    "time": "14:00 - 14:30 (30 giây)",
                    "script": (
                        "Đề tài mang lại 4 đóng góp kỹ thuật then chốt:\n"
                        "1. Kiến trúc Lookahead thời gian thực loại bỏ độ trễ tương tác.\n"
                        "2. Cơ chế Grounded RAG 3 cổng triệt tiêu hoàn toàn ảo giác của LLM.\n"
                        "3. Pipeline Giọng nói đa tầng an toàn, đa phương thức với cơ chế Fallback bền vững.\n"
                        "4. Hệ thống Quản trị tri thức Provenance chuẩn hóa theo STAR/BARS.\n\n"
                        "Về mặt giới hạn: Đề tài trung thực thừa nhận hệ thống hiện mới chứng minh được tính nhất quán và vững chắc về mặt kỹ thuật, "
                        "chưa khẳng định hiệu lực đo lường tuyển dụng (Validity/Fairness) trên quy mô xã hội rộng lớn, "
                        "và cần thêm sự thẩm định của các chuyên gia đầu ngành (SMEs) trước khi triển khai đại trà."
                    ),
                    "note": "Sự trung thực khoa học về giới hạn nghiên cứu sẽ ghi điểm rất cao với Hội đồng."
                },
                {
                    "slide_num": "Slide 31 - 32",
                    "title": "Kết luận & Lời cảm ơn",
                    "time": "14:30 - 15:00 (30 giây)",
                    "script": (
                        "Tóm lại, InterV đã hiện thực hóa thành công một giải pháp công nghệ toàn diện, tin cậy và có khả năng ứng dụng thực tiễn cao "
                        "cho bài toán phỏng vấn tự động hóa có kiểm soát.\n\n"
                        "Em xin chân thành cảm ơn Thầy ThS. Đặng Văn Lực đã tận tình hướng dẫn, "
                        "và xin trân trọng cảm ơn quý Thầy Cô trong Hội đồng đã chú ý lắng nghe.\n\n"
                        "Em xin phép được tiếp nhận các câu hỏi nhận xét và đóng góp quý báu từ Hội đồng!"
                    ),
                    "note": "Cúi đầu chào trang trọng, mỉm cười nhẹ, đứng thẳng đón nhận câu hỏi đầu tiên của Hội đồng."
                }
            ]
        }
    ]

    for part_idx, part in enumerate(slides_script):
        add_section_heading(part["part"], level=2)
        for s in part["slides"]:
            h_slide = add_section_heading(f"▶ {s['slide_num']}: {s['title']} [{s['time']}]", level=3)
            
            p_scr = doc.add_paragraph()
            p_scr.paragraph_format.line_spacing = 1.2
            p_scr.paragraph_format.space_after = Pt(4)
            r_head = p_scr.add_run("Lời nói: ")
            r_head.font.bold = True
            r_head.font.color.rgb = RGBColor(30, 58, 138)
            
            r_txt = p_scr.add_run(f"\"{s['script']}\"")
            r_txt.font.size = Pt(11)
            
            add_callout(s['note'], label="Ghi chú trình bày & Cử chỉ", bg_color="F0FDF4", border_color="16A34A")

    # --- PHẦN 3: PHÂN TÍCH CHUYÊN SÂU PIPELINE CỐT LÕI (DÀNH CHO PHẢN BIỆN) ---
    add_section_heading("PHẦN III. BẢNG TỔNG HỢP KIẾN TRÚC & PIPELINE CỐT LÕI ĐỂ PHẢN BIỆN", level=1)
    
    doc.add_paragraph(
        "Phần này tổng hợp chi tiết các tham số kỹ thuật, công nghệ và cơ chế phòng vệ của InterV "
        "để sinh viên nắm chắc và trả lời xuất sắc mọi câu hỏi đào sâu từ Hội đồng:"
    )

    deep_table = doc.add_table(rows=1, cols=3)
    deep_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    deep_table.autofit = False
    
    d_headers = ["Thành phần Pipeline", "Công nghệ & Triển khai", "Cơ chế Giải quyết Bài toán Kỹ thuật"]
    d_widths = [Inches(1.5), Inches(2.3), Inches(2.8)]
    
    d_hdr_cells = deep_table.rows[0].cells
    for i, h in enumerate(d_headers):
        d_hdr_cells[i].width = d_widths[i]
        set_cell_background(d_hdr_cells[i], "1E3A8A")
        set_cell_margins(d_hdr_cells[i], top=100, bottom=100, left=120, right=120)
        p = d_hdr_cells[i].paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(255, 255, 255)
        
    deep_data = [
        ("Lookahead Adaptive Engine", 
         "• Next.js after() Async Hook\n• Python gRPC SubmitAnswer\n• DeepSeek-Chat (V3/R1)\n• Pre-warmed TTS Buffer", 
         "• Triệt tiêu hoàn toàn độ trễ 3-5s của LLM.\n• Trả ngay Q_(i+1) (0ms) trong khi sinh câu hỏi thích ứng Q_(i+2) trong nền.\n• Fallback tức thì sang baseline question nếu mạng chậm."),
        
        ("Multimodal Speech Pipeline", 
         "• Browser AudioWorklet (16kHz PCM)\n• AssemblyAI Streaming WebSocket\n• Faster-Whisper Cuda/CPU Fallback\n• Edge/Vbee TTS + Pronunciation Lib", 
         "• Cấp temporary token ngắn hạn cho browser (bảo mật tuyệt đối).\n• Tự động kích hoạt Faster-Whisper cục bộ khi streaming lỗi/rỗng.\n• Bộ từ điển phiên âm thuật ngữ IT (React, Figma, K8s, CI/CD) giúp TTS tự nhiên."),
        
        ("Hybrid RAG & Grounding Gate", 
         "• Qdrant Vector Database\n• FastEmbed Multilingual MPNet (Dense)\n• BM25 Sparse Lexical + RRF\n• 86 Rule Markdown Files", 
         "• Kết hợp ngữ nghĩa sâu (Dense) và từ khóa chính xác (BM25).\n• Grounding 3 Cổng: Pre-context, In-generation JSON schema, Post-citation check.\n• Bất biến: Từ chối trích dẫn ngoài allow-list, tự động kích hoạt 1-shot repair."),
        
        ("SenseVoice Behavioral Analysis", 
         "• SenseVoice Small ONNX/PyTorch\n• SER (Emotion), LID (Language), AED (Events)\n• Signal Processing (RMS, dB stddev, Pause)", 
         "• Trích xuất định lượng: WPM, Pause ratio, Volume stability, Filler words.\n• Invariant an toàn: Observation-only (chỉ coaching, không chẩn đoán tâm lý), Default Neutral 50 nếu thiếu tín hiệu."),
        
        ("Grounded Rubric Evaluation", 
         "• DeepSeek Structured JSON Evaluation\n• 7 Thang điểm (STAR/BARS)\n• Exact Answer Excerpt Substring Matching\n• Mongo Distributed Transaction & Lease", 
         "• Điểm > 40 bắt buộc trích dẫn chính xác (exact excerpt) từ transcript của ứng viên.\n• Khóa nguyên tử Evaluation Lease (370s) chống race condition và trừ credit kép.")
    ]
    
    for row_idx, data in enumerate(deep_data):
        row = deep_table.add_row()
        bg_col = "F8FAFC" if row_idx % 2 == 0 else "FFFFFF"
        for col_idx, text in enumerate(data):
            cell = row.cells[col_idx]
            cell.width = d_widths[col_idx]
            set_cell_background(cell, bg_col)
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(text)
            r.font.size = Pt(9.5)
            if col_idx == 0:
                r.font.bold = True
                r.font.color.rgb = RGBColor(30, 58, 138)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # --- PHẦN 4: BỘ CÂU HỎI PHẢN BIỆN DỰ PHÒNG & KỊCH BẢN TRẢ LỜI ---
    add_section_heading("PHẦN IV. BỘ CÂU HỎI PHẢN BIỆN DỰ PHÒNG CỦA HỘI ĐỒNG (Q&A CHEAT SHEET)", level=1)

    qa_list = [
        {
            "q": "Câu 1: Tại sao nhóm không để Trình duyệt gọi trực tiếp Python Backend hay các dịch vụ AI như DeepSeek, AssemblyAI mà phải đi qua Next.js BFF và gRPC?",
            "a": (
                "Trả lời: Kính thưa Thầy Cô, đây là quyết định kiến trúc nhằm đảm bảo 3 nguyên tắc sống còn:\n"
                "1. Bảo mật Credential: Nếu Frontend gọi trực tiếp DeepSeek hay AssemblyAI, API Key sẽ bị lộ trên client. "
                "Next.js BFF đóng vai trò bảo vệ bí mật, chỉ cấp temporary token ngắn hạn cho WebSocket audio.\n"
                "2. Xác thực và Phân quyền (RBAC): Next.js kiểm tra session, quyền hạn (user/recruiter/admin), số dư credit và rate-limit "
                "trước khi ủy quyền cho Python backend xử lý.\n"
                "3. Hiệu năng & Type Safety: Giao tiếp giữa Next.js và Python sử dụng gRPC (HTTP/2 binary Protobuf), "
                "cho tốc độ truyền tải cực nhanh, độ trễ thấp hơn nhiều so với REST JSON nội bộ, và ép buộc tính toàn vẹn kiểu dữ liệu giữa TypeScript và Python."
            )
        },
        {
            "q": "Câu 2: Nếu ứng viên cố tình trả lời lạc đề hoặc trả lời quá ngắn (chỉ 1-2 từ), hệ thống xử lý như thế nào?",
            "a": (
                "Trả lời: Kính thưa Thầy Cô, InterV xử lý tình huống này ở cả 2 tầng:\n"
                "1. Trong lúc phỏng vấn (Lookahead Probing): DeepSeek nhận toàn bộ lịch sử QA. Khi phát hiện câu trả lời thiếu thông tin (gap), "
                "DeepSeek sẽ sinh câu hỏi thích ứng tiếp theo để đào sâu (follow-up probe) nhằm yêu cầu ứng viên làm rõ hoặc đưa ra ví dụ cụ thể.\n"
                "2. Trong lúc chấm điểm (Exact Answer Excerpt Check): Hệ thống áp dụng quy tắc bắt buộc: Điểm số trên 40 phải có đoạn trích dẫn chứng minh từ transcript. "
                "Nếu câu trả lời quá ngắn hoặc lạc đề, không trích xuất được evidence hợp lệ, DeepSeek sẽ chấm điểm thấp và chỉ ra lỗi cụ thể trong phần 'mistakes' và 'feedback'."
            )
        },
        {
            "q": "Câu 3: Làm sao đảm bảo SenseVoice không đưa ra các nhận xét mang tính kỳ thị hoặc sai lệch tâm lý của ứng viên?",
            "a": (
                "Trả lời: Kính thưa Thầy Cô, hệ thống cài đặt Bất biến an toàn (Safety Invariant) ở mức mã nguồn:\n"
                "1. Nguyên tắc Observation-only: SenseVoice chỉ đo lường các thuộc tính vật lý âm thanh (Tốc độ nói WPM, độ biến thiên âm lượng dB, tỷ lệ ngập ngừng, từ đệm) "
                "và cảm xúc bề mặt (SER), hoàn toàn KHÔNG đưa ra kết luận về tính cách, sự trung thực hay bệnh lý tâm thần.\n"
                "2. Giá trị mặc định trung tính (Neutral Default = 50): Nếu audio nhiễu, đứt quãng hoặc không đủ thời lượng để phân tích, "
                "hệ thống tự động gán điểm trung tính 50, không bao giờ trừ điểm hay đánh trượt ứng viên dựa trên tín hiệu âm thanh."
            )
        },
        {
            "q": "Câu 4: Cơ chế Grounding Allow-list hoạt động thế nào khi DeepSeek bị 'ảo giác' sinh ra một trích dẫn không có thật?",
            "a": (
                "Trả lời: Kính thưa Thầy Cô, khi Python backend chuẩn bị ngữ cảnh cho DeepSeek, nó đóng gói một tập hợp Evidence IDs hợp lệ (Grounding Package). "
                "Khi DeepSeek trả về JSON kết quả, hàm `validate_grounding_ids()` sẽ đối chiếu từng ID trong citation với Allow-list. "
                "Nếu xuất hiện ID không hợp lệ, hệ thống sẽ chặn đứng dữ liệu tại Citation Gate, "
                "đồng thời kích hoạt cơ chế 1-shot repair: Gửi lại prompt nhắc nhở kèm đúng danh sách Allow-list để DeepSeek tự động sửa và trả về kết quả chuẩn xác."
            )
        }
    ]

    for item in qa_list:
        p_q = doc.add_paragraph()
        p_q.paragraph_format.space_before = Pt(8)
        p_q.paragraph_format.space_after = Pt(2)
        r_q = p_q.add_run(f"❓ {item['q']}")
        r_q.font.bold = True
        r_q.font.color.rgb = RGBColor(185, 28, 28)
        
        p_a = doc.add_paragraph()
        p_a.paragraph_format.space_after = Pt(8)
        r_a = p_a.add_run(item['a'])
        r_a.font.size = Pt(10.5)

    doc_path = r"d:\project\InterV\Review\Kich_Ban_Thuyet_Trinh_KLTN_InterV_15Phut.docx"
    doc.save(doc_path)
    print(f"Document successfully created and saved at: {doc_path}")

if __name__ == "__main__":
    create_document()
